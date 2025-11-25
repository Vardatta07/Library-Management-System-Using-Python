from docx import Document
from docx.shared import Pt

# Create a new Document
doc = Document()

# Title
doc.add_heading('Cybersecurity Report', 0)

# Executive Summary
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    "This report reviews common cybersecurity threats and explains strategies and best practices to mitigate risks. "
    "It covers malware, ransomware, phishing, DDoS, insider threats, zero-day attacks, and supply chain compromises."
)

# Introduction
doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    "Cybersecurity protects digital systems from attacks. This report outlines major threats, mitigation options, "
    "and long‑term security practices used in modern environments."
)

# Common Cybersecurity Threats
doc.add_heading('Common Cybersecurity Threats', level=1)
threats = [
    "1. Malware – Viruses, worms, trojans, spyware.",
    "2. Ransomware – Encrypts data and demands payment.",
    "3. Phishing – Trick users into leaking information.",
    "4. DDoS – Overloads networks.",
    "5. Insider Threats – Misuse of internal access.",
    "6. Zero‑Day Attacks – Exploit unknown flaws.",
    "7. Supply Chain Attacks – Compromise trusted vendors."
]
for threat in threats:
    doc.add_paragraph(threat)

# Mitigation Strategies
doc.add_heading('Mitigation Strategies', level=1)
strategies = [
    "Use antivirus, EDR, and regular patching.",
    "Maintain offline backups.",
    "Enable MFA and email filtering.",
    "Use DDoS protection and network segmentation.",
    "Implement RBAC, monitoring, and audits.",
    "Apply defense‑in‑depth and virtual patching.",
    "Evaluate software vendors and check update integrity."
]
for strategy in strategies:
    p = doc.add_paragraph(strategy, style='List Bullet')

# Best Practices
doc.add_heading('Best Practices', level=1)
practices = [
    "Firewalls, segmentation, secure VPN.",
    "Encrypt data in transit and at rest.",
    "Least‑privilege access and MFA.",
    "System hardening and timely updates.",
    "SIEM monitoring and strong incident response."
]
for practice in practices:
    p = doc.add_paragraph(practice, style='List Bullet')

# Conclusion
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    "Cybersecurity threats are constantly evolving, but strong protection is possible using layered defenses, "
    "monitoring, good configuration, and trained users."
)

# References
doc.add_heading('References', level=1)
references = [
    "NIST Cybersecurity Framework",
    "OWASP Guidelines",
    "ISO 27001 Standard",
    "CERT Advisories"
]
for ref in references:
    p = doc.add_paragraph(ref, style='List Bullet')

# Save the document
file_path = '/mnt/data/Cybersecurity_Report.docx'
doc.save(file_path)

file_path